import os
import sys
import requests
import unicodedata
import time
import json
import shutil
import threading
import subprocess
import re
from datetime import timedelta, datetime
from pathlib import Path

# Check and install required dependencies
def check_dependencies():
    """Verify and install required dependencies"""
    try:
        # First check if the packages are installed
        import importlib.util
        nltk_spec = importlib.util.find_spec("nltk")
        docx_spec = importlib.util.find_spec("docx")  # python-docx imports as docx
        lxml_spec = importlib.util.find_spec("lxml")
        
        if all([nltk_spec, docx_spec, lxml_spec]):
            # Then try to import them to ensure they work
            import nltk
            import docx
            import lxml.etree
            print("Required dependencies are already installed.")
            return True
        else:
            raise ImportError("One or more required packages not found")
    except ImportError:
        print("Installing required dependencies...")
        try:
            import pip
            pip.main(['install', 'nltk', 'python-docx', 'lxml'])
            # After installing, try importing them again
            import nltk
            import docx
            import lxml.etree
            print("Successfully installed dependencies.")
            return True
        except Exception as e:
            print(f"Failed to install dependencies: {e}")
            print("Please run: pip install nltk python-docx lxml")
            return False

# Check dependencies first
if not check_dependencies():
    sys.exit(1)

# Now we can safely import these
import nltk
import docx
from lxml import etree

# Download NLTK data
try:
    nltk.download("punkt", quiet=True)
except Exception as e:
    print(f"Warning: Failed to download NLTK data: {e}")
    print("Some sentence splitting may not work correctly.")

# Default HuggingFace API endpoints
HF_API_URL = os.getenv("HF_API_URL", "https://bz2eki98bvwdoh9l.us-east4.gcp.endpoints.huggingface.cloud")
HF_MODEL = os.getenv("HF_MODEL", "mistralai/Mistral-7B-Instruct-v0.2") # Default to Mistral model if not specified
HF_TOKEN = os.getenv("HF_TOKEN")
if not HF_TOKEN:
    print("Please set the HF_TOKEN environment variable")
    sys.exit(1)

# Translation template environment variables
TRANSLATE_TEMPLATE_En = os.getenv("TRANSLATE_TEMPLATE_En", "")
TRANSLATE_TEMPLATE_Ja = os.getenv("TRANSLATE_TEMPLATE_Ja", "")

def check_translation_templates():
    """Check for translation templates in environment variables and set up defaults if needed"""
    global TRANSLATE_TEMPLATE_En, TRANSLATE_TEMPLATE_Ja
    
    # Check if environment variable is set
    if not TRANSLATE_TEMPLATE_En:
        print("\nNOTE: TRANSLATE_TEMPLATE_En environment variable not set.")
        print("For best results, set this environment variable with a template like:")
        print("-------------------------------------------------------------------")
        print('TRANSLATE_TEMPLATE_En="[INST] You are a professional Japanese translator. Your task is to translate the following English text into natural, fluent Japanese.\n\nTranslation style: Young adult fantasy novel for vertical writing (tategaki).\nExpress subtle wonder, emotional depth, and a slightly melancholic tone.\nUse appropriate Japanese punctuation (「」for quotes, 。for periods, etc).\nUse natural Japanese expressions rather than literal translations.\n\nIMPORTANT: Return ONLY the Japanese translation. No explanations or English text.\n\nText to translate: {text} [/INST]"')
        print("-------------------------------------------------------------------")
        print("Using a default template for now. Set the environment variable for best results.\n")
    else:
        print("Found TRANSLATE_TEMPLATE_En in environment variables.")
        
    # Check template format
    if TRANSLATE_TEMPLATE_En and "{text}" not in TRANSLATE_TEMPLATE_En:
        print("WARNING: TRANSLATE_TEMPLATE_En must contain {text} placeholder. Using default template.")
        TRANSLATE_TEMPLATE_En = ""
        
    return True

# Check translation templates
check_translation_templates()

# Allow fallback to other models if the first one fails
BACKUP_MODELS = [
    "meta-llama/Llama-2-70b-chat-hf",
    "gemini/gemini-pro",
    "google/gemma-7b-it",
    "mistralai/Mixtral-8x7B-Instruct-v0.1"
]

HEADERS = {"Authorization": f"Bearer {HF_TOKEN}"}

# Function to test the API connection
def test_api_connection():
    """Test the API connection and return diagnostic information"""
    print(f"Testing API connection to {HF_API_URL}...")
    short_prompt = "Translate this to Japanese: Hello world."
    payload = {
        "inputs": short_prompt,
        "parameters": {
            "max_new_tokens": 100,
            "return_full_text": False,
        },
    }
    
    try:
        response = requests.post(HF_API_URL, headers=HEADERS, json=payload, timeout=10)
        print(f"API Status Code: {response.status_code}")
        
        if response.status_code == 200:
            try:
                data = response.json()
                print(f"API Response structure: {json.dumps(data, indent=2)[:500]}...")
                return True
            except json.JSONDecodeError:
                print(f"Could not parse API response as JSON. Response text: {response.text[:500]}...")
                return False
        elif response.status_code == 401 or response.status_code == 403:
            print("Authentication error. Please check your HF_TOKEN.")
            return False
        elif response.status_code == 404:
            print("API endpoint not found. Please verify the API URL.")
            return False
        else:
            print(f"Unexpected status code: {response.status_code}, Response: {response.text[:500]}...")
            return False
    except requests.exceptions.RequestException as e:
        print(f"Connection error: {e}")
        return False

def test_translation_template():
    """Test the translation with the current template"""
    global TRANSLATE_TEMPLATE_En
    
    print("\nTesting translation with current template...")
    print(f"Using API endpoint: {HF_API_URL}")
    print(f"Using model: {HF_MODEL}")
    
    # Simple test phrase
    test_text = "The old maple tree whispered secrets as the autumn breeze passed through its branches."
    
    try:
        # Translate the test phrase
        translated = translate_text(test_text)
        
        # Check if translation succeeded
        has_jp_chars = any(0x3000 <= ord(char) <= 0x9FFF for char in translated)
        
        if has_jp_chars:
            print("\n✓ Translation test successful!")
            print(f"Input:  {test_text}")
            print(f"Output: {translated}")
            return True
        else:
            print("\n✗ Translation test failed. No Japanese characters in output.")
            print(f"Output received: {translated}")
            return False
    except Exception as e:
        print(f"\n✗ Translation test error: {e}")
        return False

OUTPUT_FILE = "JP.txt"
CHECKPOINT_FILE = "translation_checkpoint.json"
BACKUP_DIR = "translate"
GCS_BUCKET = "gs://cdn-othertales-co/"
BACKUP_INTERVAL = 120  # Backup every 2 minutes

# Create backup directory if it doesn't exist
if not os.path.exists(BACKUP_DIR):
    os.makedirs(BACKUP_DIR, exist_ok=True)

# Flag to control the backup thread
backup_thread_active = False

def cloud_backup_thread():
    """Background thread that backs up the translation file to GCS every 2 minutes"""
    global backup_thread_active
    
    while backup_thread_active:
        try:
            if os.path.exists(OUTPUT_FILE):
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                
                # Create backup filename with timestamp and original filename
                filename = Path(OUTPUT_FILE).name
                base_name = Path(filename).stem
                extension = Path(filename).suffix
                timestamped_filename = f"{base_name}_{timestamp}{extension}"
                
                # Backup to the main file (overwrite) and also to a timestamped version
                print(f"[{timestamp}] Backing up {OUTPUT_FILE} to {GCS_BUCKET}")
                
                # Execute gsutil command to copy the file to Google Cloud Storage
                # First, copy to the main filename (overwriting previous version)
                result = subprocess.run(
                    ["gsutil", "cp", OUTPUT_FILE, GCS_BUCKET],
                    capture_output=True, 
                    text=True
                )
                
                # Then, copy to a timestamped filename (for versioning)
                timestamped_result = subprocess.run(
                    ["gsutil", "cp", OUTPUT_FILE, f"{GCS_BUCKET}{timestamped_filename}"],
                    capture_output=True, 
                    text=True
                )
                
                if result.returncode == 0 and timestamped_result.returncode == 0:
                    print(f"[{timestamp}] Backup successful")
                else:
                    if result.returncode != 0:
                        print(f"[{timestamp}] Main backup failed: {result.stderr}")
                    if timestamped_result.returncode != 0:
                        print(f"[{timestamp}] Timestamped backup failed: {timestamped_result.stderr}")
        except Exception as e:
            print(f"Error during cloud backup: {e}")
        
        # Sleep for the specified interval
        time.sleep(BACKUP_INTERVAL)

def start_backup_thread():
    """Start the background backup thread"""
    global backup_thread_active
    
    # Check if gsutil is available
    try:
        result = subprocess.run(
            ["gsutil", "version"], 
            capture_output=True,
            text=True
        )
        if result.returncode != 0:
            print("Warning: gsutil not found or not configured. Cloud backups will be disabled.")
            print("Please install and configure Google Cloud SDK if you want cloud backups.")
            return False
    except FileNotFoundError:
        print("Warning: gsutil not found. Cloud backups will be disabled.")
        print("Please install Google Cloud SDK if you want cloud backups.")
        return False
    
    backup_thread_active = True
    backup_thread = threading.Thread(target=cloud_backup_thread, daemon=True)
    backup_thread.start()
    print("Cloud backup thread started. Will backup every 2 minutes.")
    return True

def stop_backup_thread():
    """Stop the background backup thread"""
    global backup_thread_active
    backup_thread_active = False
    print("Cloud backup thread stopped.")

def normalize_punctuation(text):
    # First apply NFKC normalization
    text = unicodedata.normalize("NFKC", text)
    
    # Replace Western punctuation with Japanese equivalents suitable for Tategaki
    replacements = {
        '.': '。',    # Western period to Japanese full stop
        ',': '、',    # Western comma to Japanese comma
        '!': '！',    # Exclamation mark to full-width
        '?': '？',    # Question mark to full-width
        '(': '（',    # Opening parenthesis to full-width
        ')': '）',    # Closing parenthesis to full-width
        '"': '「',    # Opening quote to Japanese opening quote
        '"': '」',    # Closing quote to Japanese closing quote
        "'": '『',    # Opening single quote to Japanese nested opening quote
        "'": '』',    # Closing single quote to Japanese nested closing quote
        ':': '：',    # Colon to full-width
        ';': '；',    # Semicolon to full-width
        '-': 'ー',    # Hyphen to Japanese long vowel mark
        '...': '…',   # Ellipsis to Japanese ellipsis
        '–': '—',     # En dash to Em dash
    }
    
    # Apply replacements
    for western, japanese in replacements.items():
        text = text.replace(western, japanese)
    
    # Handle special cases for quotes that may be ASCII quotes
    text = text.replace('"', '「').replace('"', '」')
    text = text.replace("'", '『').replace("'", '』')
    
    # Additional specific replacements for common patterns
    text = text.replace('--', '—')
    
    return text

def clean_translation_output(text):
    """
    Clean up the translation output to extract only the Japanese text.
    Removes formatting markers, prefixes, and other non-translation content.
    
    More robust handling of different response types:
    1. Structured responses with markers (e.g. "**Japanese Translation:**")
    2. Mixed language responses (extract only Japanese parts)
    3. Pure English responses (return error message)
    """
    # Helper function to check if a character is Japanese
    def is_japanese_char(char):
        code = ord(char)
        return (0x3000 <= code <= 0x9FFF or   # CJK unified ideographs, Hiragana, Katakana
                0xF900 <= code <= 0xFAFF or   # CJK compatibility ideographs
                0xFF00 <= code <= 0xFFEF)      # Half-width and full-width forms
    
    # Helper function to check if a character is Japanese punctuation
    def is_japanese_punct(char):
        return char in "「」『』（）！？。、：；"
    
    # First, check if there are any Japanese characters in the response
    has_japanese = any(is_japanese_char(char) for char in text)
    
    if not has_japanese:
        print(f"WARNING: No Japanese characters found in the response: {text[:100]}...")
        return "［翻訳エラー］"  # Japanese for [translation error]
    
    # Calculate the Japanese character ratio
    jp_char_count = sum(1 for char in text if is_japanese_char(char) or is_japanese_punct(char))
    total_chars = len(text.strip())
    jp_ratio = jp_char_count / total_chars if total_chars > 0 else 0
    
    # Very clean response (mostly Japanese) - just do basic cleanup
    if jp_ratio > 0.7:  # If more than 70% is Japanese, it's likely a clean response
        # Just strip and remove unnecessary quotes
        cleaned_text = text.strip()
        if (cleaned_text.startswith('"') and cleaned_text.endswith('"')) or \
           (cleaned_text.startswith("'") and cleaned_text.endswith("'")):
            cleaned_text = cleaned_text[1:-1].strip()
        return cleaned_text
    
    # Common patterns to remove from model output
    patterns = [
        # Match anything with "Original Sentence:" or "Japanese Translation:" markers
        r'\*\*Original Sentence[：:]\*\*.*?(?=\*\*Japanese Translation[：:]\*\*|\Z)',
        r'\*\*Japanese Translation[：:]\*\*\s*',
        r'\*\*[^*]+?\*\*',  # Any **text** format markers
        r'##\s*Japanese\s*Translation[：:]\s*',  # Section headers
        r'##\s*[^#]+',  # Any ## header format
        r'Japanese Translation[：:]',  # Plain text markers
        r'Translation[：:]',  # More general translation markers
        r'.*?(?=日本語|[ぁ-んァ-ン])',  # Match everything up to first Japanese character
    ]
    
    # Check if the text contains markers that indicate a structured response
    has_markers = any(re.search(pattern, text) for pattern in [
        r'\*\*Japanese Translation', 
        r'##\s*Japanese\s*Translation',
        r'Japanese Translation[：:]',
        r'Translation[：:]'
    ])
    
    # Try to extract Japanese text using different methods
    extracted_text = None
    
    # 1. MARKER-BASED EXTRACTION
    if has_markers:
        cleaned_text = text
        for pattern in patterns:
            cleaned_text = re.sub(pattern, '', cleaned_text, flags=re.DOTALL)
        
        # If we got good Japanese content, use it
        jp_char_count = sum(1 for char in cleaned_text if is_japanese_char(char))
        if jp_char_count > 10:  # Arbitrary threshold for "enough" Japanese text
            extracted_text = cleaned_text
    
    # 2. SENTENCE-BASED EXTRACTION
    if not extracted_text:
        # Find continuous Japanese segments with Japanese punctuation
        japanese_segments = re.findall(r'[「」『』（）！？。、]?[^\n\r]+?[「」『』（）！？。、]', text)
        
        if japanese_segments:
            # Filter segments to only those with Japanese characters
            japanese_segments = [
                seg for seg in japanese_segments 
                if any(is_japanese_char(char) for char in seg)
            ]
            
            if japanese_segments:
                extracted_text = "".join(japanese_segments)
    
    # 3. LINE-BASED EXTRACTION
    if not extracted_text:
        lines = text.split('\n')
        japanese_lines = []
        
        for line in lines:
            # Count Japanese characters in this line
            jp_chars = sum(1 for c in line if is_japanese_char(c))
            
            # If there are Japanese characters and the line is mostly Japanese or contains sentences
            if jp_chars > 0 and ('。' in line or '！' in line or '？' in line or jp_chars / len(line) > 0.3):
                # Extract just Japanese segments within this line
                jp_segments = re.findall(r'[「」『』（）！？。、]?[^a-zA-Z\(\)]+[「」『』（）！？。、]?', line)
                if jp_segments:
                    japanese_lines.append("".join(jp_segments))
                else:
                    japanese_lines.append(line)
        
        if japanese_lines:
            extracted_text = '\n'.join(japanese_lines)
    
    # 4. CHARACTER-BY-CHARACTER EXTRACTION (last resort)
    if not extracted_text:
        # Extract continuous Japanese characters sequences
        japanese_parts = []
        current_part = []
        
        for char in text:
            if is_japanese_char(char) or is_japanese_punct(char) or char.isspace():
                current_part.append(char)
            elif current_part:  # End of a Japanese segment
                segment = ''.join(current_part).strip()
                if any(is_japanese_char(c) for c in segment):
                    japanese_parts.append(segment)
                current_part = []
        
        # Don't forget the last segment
        if current_part:
            segment = ''.join(current_part).strip()
            if any(is_japanese_char(c) for c in segment):
                japanese_parts.append(segment)
        
        if japanese_parts:
            extracted_text = ' '.join(japanese_parts)
    
    # Use the extracted text if we found any, otherwise use the original
    final_text = extracted_text if extracted_text else text
    
    # Final cleanup
    final_text = final_text.strip()
    
    # Remove any remaining English phrases in parentheses - common in translations
    final_text = re.sub(r'\([^)]*[a-zA-Z][^)]*\)', '', final_text)
    
    # If final text is still wrapped in quotes, remove them
    if (final_text.startswith('"') and final_text.endswith('"')) or \
       (final_text.startswith("'") and final_text.endswith("'")):
        final_text = final_text[1:-1].strip()
    
    # If we don't have enough Japanese characters in the final result, return error
    jp_char_count = sum(1 for char in final_text if is_japanese_char(char))
    if jp_char_count < 5:  # Arbitrary threshold for "enough" Japanese text
        print(f"WARNING: Too few Japanese characters in cleaned output: {final_text}")
        return "［翻訳エラー］"  # Japanese for [translation error]
    
    return final_text

def build_translation_prompt(text_to_translate, model_name=None):
    """Build an appropriate translation prompt based on model type"""
    global TRANSLATE_TEMPLATE_En, TRANSLATE_TEMPLATE_Ja
    
    # If environment variable template is available, use it
    if TRANSLATE_TEMPLATE_En:
        # Replace {text} with the actual text to translate
        return TRANSLATE_TEMPLATE_En.replace("{text}", text_to_translate)
    # If Japanese template is available and we detect Japanese input
    elif TRANSLATE_TEMPLATE_Ja and any(0x3000 <= ord(char) <= 0x9FFF for char in text_to_translate):
        return TRANSLATE_TEMPLATE_Ja.replace("{text}", text_to_translate)
    
    # Otherwise, use model-specific templates
    # For Mistral and similar instruction models
    elif model_name and ("mistral" in model_name.lower() or "llama" in model_name.lower() or 
                       "gemma" in model_name.lower() or "mixtral" in model_name.lower()):
        return (
            "[INST] You are a professional Japanese translator. Translate this English text to Japanese:\n\n"
            "Translation style: Young adult fantasy novel for vertical writing (tategaki). "
            "Express subtle wonder, emotional depth, and a slightly melancholic tone.\n"
            "IMPORTANT: Output Japanese text ONLY. No explanations or English text.\n\n"
            f"Text to translate: {text_to_translate} [/INST]"
        )
    # For Claude-like models
    elif model_name and "claude" in model_name.lower():
        return (
            "Human: You are a professional Japanese translator. Translate this English text to Japanese:\n\n"
            "Translation style: Young adult fantasy novel for vertical writing (tategaki). "
            "Express subtle wonder, emotional depth, and a slightly melancholic tone.\n"
            "IMPORTANT: Output Japanese text ONLY. No explanations or English text.\n\n"
            f"Text to translate: {text_to_translate}\n\n"
            "Assistant: "
        )
    # For GPT-like models or default
    else:
        return (
            "あなたは優秀な日本語翻訳者です。次の英語の文章を日本語に翻訳してください。\n\n"
            "翻訳スタイルは日本の若者向けファンタジー小説（縦書き用）です。微妙な不思議さ、感情の深さ、"
            "少し風変わりで物悲しい調子を表現しつつ、そのジャンルのために書かれたかのように自然に言葉が流れるようにしてください。\n\n"
            "回答は日本語の翻訳のみを提供してください。フォーマット、説明、英語のテキストは含めないでください。\n\n"
            "翻訳する文章：\n" + text_to_translate + "\n\n"
            "日本語翻訳："
        )

def translate_text(text_to_translate, max_retries=3, retry_delay=5):
    """Translate text with retry logic for API failures and model fallback"""
    global HF_API_URL, HF_MODEL, TRANSLATE_TEMPLATE_En
    
    # Try with primary model first
    current_model = HF_MODEL
    models_to_try = [current_model] + BACKUP_MODELS
    
    # Check if environment variable template is available - if not, set a recommended one
    if not TRANSLATE_TEMPLATE_En:
        print("Note: TRANSLATE_TEMPLATE_En environment variable not set. Using default template.")
        # You may want to suggest setting this environment variable
        recommended_template = (
            "[INST] You are a professional Japanese translator. Your task is to translate the following English text into natural, fluent Japanese.\n\n"
            "Translation style: Young adult fantasy novel for vertical writing (tategaki).\n"
            "Express subtle wonder, emotional depth, and a slightly melancholic tone.\n"
            "Use appropriate Japanese punctuation (「」for quotes, 。for periods, etc).\n"
            "Use natural Japanese expressions rather than literal translations.\n\n"
            "IMPORTANT: Return ONLY the Japanese translation. No explanations or English text.\n\n"
            "Text to translate: {text} [/INST]"
        )
        # We'll use this for now but won't permanently set the env var
        TRANSLATE_TEMPLATE_En = recommended_template
    
    for model_index, model in enumerate(models_to_try):
        # Build appropriate prompt for the model
        prompt = build_translation_prompt(text_to_translate, model)
        
        payload = {
            "inputs": prompt,
            "parameters": {
                "max_new_tokens": 1024,
                "return_full_text": False,
                "temperature": 0.7,
                "top_p": 0.95,
            },
        }
        
        # Add model parameter if needed (some API endpoints require it)
        if model_index > 0:  # For backup models
            payload["model"] = model
            print(f"Trying with backup model: {model}")
        
        for attempt in range(max_retries):
            try:
                # Print diagnostic info for the first attempt
                if attempt == 0 and model_index == 0:
                    print(f"Sending request to: {HF_API_URL}")
                    print(f"With headers: {HEADERS}")
                    if model_index > 0:
                        print(f"Using model: {model}")
                
                response = requests.post(HF_API_URL, headers=HEADERS, json=payload, timeout=30)
                
                # Handle various response codes
                if response.status_code == 429:  # Rate limit
                    wait_time = retry_delay * (attempt + 1)  # Exponential backoff
                    print(f"Rate limited. Waiting {wait_time} seconds before retrying...")
                    time.sleep(wait_time)
                    continue
                    
                if response.status_code != 200:
                    print(f"API Error {response.status_code}: {response.text[:500]}...")
                    if attempt == max_retries - 1 and model_index < len(models_to_try) - 1:
                        # We'll try another model
                        break
                    elif attempt < max_retries - 1:
                        # We'll retry with same model
                        wait_time = retry_delay * (attempt + 1)
                        print(f"Retrying in {wait_time} seconds...")
                        time.sleep(wait_time)
                        continue
                    else:
                        # Last model, last attempt
                        raise Exception(f"All models failed. Last error: API Error {response.status_code}: {response.text}")
                
                # Parse the response
                try:
                    data = response.json()
                    # Debug: print raw response structure
                    if attempt == 0 and model_index == 0:
                        print(f"Raw API response structure: {type(data)}")
                        if isinstance(data, list) and len(data) > 0:
                            print(f"First element type: {type(data[0])}")
                            if hasattr(data[0], 'keys'):
                                print(f"Keys: {list(data[0].keys())}")
                    
                    # Extract the text, handling different response formats
                    if isinstance(data, list) and len(data) > 0 and isinstance(data[0], dict) and "generated_text" in data[0]:
                        jp_text = data[0]["generated_text"].strip()
                    elif isinstance(data, dict) and "choices" in data and len(data["choices"]) > 0:
                        # Handle OpenAI-like format
                        jp_text = data["choices"][0]["message"]["content"].strip()
                    elif isinstance(data, dict) and "generated_text" in data:
                        jp_text = data["generated_text"].strip()
                    else:
                        # If we can't extract using known formats, use the whole response
                        jp_text = str(data)
                    
                    # Clean up any formatting or markers in the translation
                    jp_text = clean_translation_output(jp_text)
                    
                    # If we got no Japanese text back after cleaning, retry
                    if jp_text == "［翻訳エラー］":
                        if attempt < max_retries - 1:
                            print(f"No Japanese text found in response, retrying with clearer instructions...")
                            # Try a different prompt format
                            prompt = (
                                "[重要] 以下の文章を日本語に翻訳してください。英語を含めないでください。" + 
                                "英語: " + text_to_translate + " 日本語: "
                            )
                            payload["inputs"] = prompt
                            continue
                        elif model_index < len(models_to_try) - 1:
                            # Try next model
                            print(f"No Japanese text found with {model}, trying next model...")
                            break
                        else:
                            # Last model, return whatever we have
                            print(f"Warning: No Japanese text found in final response: {jp_text[:100]}...")
                            return "［翻訳できませんでした］"  # Japanese for [could not translate]
                    
                    # Apply punctuation normalization and return successful translation
                    return normalize_punctuation(jp_text)
                
                except (json.JSONDecodeError, KeyError, IndexError, TypeError) as e:
                    print(f"Error parsing API response: {e}")
                    print(f"Raw response: {response.text[:500]}...")
                    
                    if attempt < max_retries - 1:
                        wait_time = retry_delay * (attempt + 1)
                        print(f"Retrying in {wait_time} seconds...")
                        time.sleep(wait_time)
                        continue
                    elif model_index < len(models_to_try) - 1:
                        # Try next model
                        break
                    else:
                        # Last model, last attempt
                        raise Exception(f"Failed to parse response from all models: {e}")
            
            except requests.exceptions.RequestException as e:
                print(f"Connection error: {e}")
                if attempt < max_retries - 1:
                    wait_time = retry_delay * (attempt + 1)
                    print(f"Retrying in {wait_time} seconds...")
                    time.sleep(wait_time)
                elif model_index < len(models_to_try) - 1:
                    # Try next model
                    print(f"Connection error with {model}, trying next model...")
                    break
                else:
                    # Last model, last attempt
                    raise Exception(f"Connection failed for all models: {e}")
    
    # If we've exhausted all models and retries
    raise Exception("Translation failed after trying all models")

def calculate_eta(start_time, current_progress, total_items):
    elapsed = time.time() - start_time
    if current_progress == 0:  # Avoid division by zero
        return "calculating..."
    
    progress_ratio = current_progress / total_items
    if progress_ratio == 0:  # Avoid division by zero
        return "calculating..."
    
    total_time_estimate = elapsed / progress_ratio
    remaining_time = total_time_estimate - elapsed
    
    # Format as HH:MM:SS
    return str(timedelta(seconds=int(remaining_time)))

def save_checkpoint(checkpoint_data):
    """Save checkpoint data to allow resuming translation"""
    with open(CHECKPOINT_FILE, 'w', encoding='utf-8') as f:
        json.dump(checkpoint_data, f)
        f.flush()
        os.fsync(f.fileno())  # Force write to disk

def create_tategaki_document(output_text_file, output_docx_file="JP_tategaki.docx"):
    """
    Create a properly formatted tategaki (vertical writing) Word document
    from the Japanese translation text file.
    
    Args:
        output_text_file: Path to the text file containing Japanese translation
        output_docx_file: Path to save the formatted Word document
    """
    try:
        import docx
        from docx.shared import Pt, Inches, Mm, RGBColor
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
        from docx.enum.section import WD_ORIENT, WD_SECTION
        import zipfile
        import os
        from lxml import etree
        
        print(f"Creating tategaki Word document: {output_docx_file}")
        
        # Read the translated content
        with open(output_text_file, 'r', encoding='utf-8') as f:
            content = f.read()
            
        # Split content by paragraphs (empty lines)
        paragraphs = [p for p in content.split('\n\n') if p.strip()]
        
        # Create new document with proper Japanese settings
        doc = docx.Document()
        
        # Set document language to Japanese
        doc.styles['Normal'].font.name = 'MS Mincho'  # Classic Japanese font
        doc.styles['Normal'].font.size = Pt(10.5)     # Standard Japanese novel size
        doc.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        doc.styles['Normal'].paragraph_format.space_after = Pt(0)
        doc.styles['Normal'].paragraph_format.space_before = Pt(0)
        
        # Set up sections for tategaki (vertical writing)
        section = doc.sections[0]
        section.orientation = WD_ORIENT.PORTRAIT
        
        # Standard Japanese novel B6 size (128mm x 182mm) - slightly smaller than JIS B6
        section.page_width = Mm(128)
        section.page_height = Mm(182)
        
        # Set appropriate margins for tategaki fiction
        section.left_margin = Mm(20)   # Top margin in tategaki
        section.right_margin = Mm(20)  # Bottom margin in tategaki
        section.top_margin = Mm(15)    # Right margin in tategaki
        section.bottom_margin = Mm(15) # Left margin in tategaki
        
        # Add title if the first paragraph appears to be a title (short paragraph)
        if paragraphs and len(paragraphs[0]) < 30:  # Typically titles are short
            title_paragraph = doc.add_paragraph()
            title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            title_run = title_paragraph.add_run(paragraphs[0])
            title_run.font.size = Pt(14)
            title_run.font.bold = True
            # Add proper spacing after title
            title_paragraph.paragraph_format.space_after = Pt(24)
            paragraphs = paragraphs[1:]  # Remove title from content to process
            
            # Add subtitle if second paragraph is also short
            if paragraphs and len(paragraphs[0]) < 40:
                subtitle_paragraph = doc.add_paragraph()
                subtitle_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                subtitle_run = subtitle_paragraph.add_run(paragraphs[0])
                subtitle_run.font.size = Pt(12)
                subtitle_run.italic = True
                subtitle_paragraph.paragraph_format.space_after = Pt(36)
                paragraphs = paragraphs[1:]
        
        # Add content paragraphs with proper Japanese formatting
        for i, para_text in enumerate(paragraphs):
            para = doc.add_paragraph()
            
            # Special formatting for chapter beginnings (typically shorter paragraphs)
            if len(para_text) < 40 and para_text.strip().startswith('第') or '章' in para_text[:10]:
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = para.add_run(para_text.strip())
                run.font.size = Pt(12)
                run.bold = True
                para.paragraph_format.space_before = Pt(24)
                para.paragraph_format.space_after = Pt(24)
            else:
                # Regular paragraph
                run = para.add_run(para_text.strip())
                para.paragraph_format.first_line_indent = Pt(10.5)  # Standard indentation for Japanese fiction
                
                # Add proper line spacing
                para.paragraph_format.line_spacing = 1.5
                
                # Add proper spacing between paragraphs, but not too much
                para.paragraph_format.space_after = Pt(12)
        
        # Save the document
        temp_docx = "temp_" + output_docx_file
        doc.save(temp_docx)
        
        # Now we need to modify the Word XML directly to enable tategaki
        # This is done by editing the document.xml file inside the .docx file (which is a zip)
        
        # Define XML namespace
        w_namespace = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        
        # Open the saved docx file as a zip
        with zipfile.ZipFile(temp_docx, 'r') as zip_ref:
            # Extract document.xml from the .docx
            zip_ref.extract('word/document.xml')
            
            # Parse the XML - fix the missing parser parameter
            parser = etree.XMLParser(recover=True)
            tree = etree.parse('word/document.xml', parser)
            root = tree.getroot()
            
            # Add vertical text orientation to the document section properties
            # Find the section properties
            for sectPr in root.findall('.//' + w_namespace + 'sectPr'):
                # Add vertical text direction (if doesn't exist)
                textDirection = sectPr.find(w_namespace + 'textDirection')
                if textDirection is None:
                    # Fix the missing nsmap parameter by using an empty dict for attrib and None for nsmap
                    textDirection = etree.SubElement(sectPr, w_namespace + 'textDirection', {}, None)
                textDirection.set(w_namespace + 'val', 'tbRl')  # Top to bottom, right to left
                
                # Set document grid for proper spacing (essential for Tategaki)
                docGrid = sectPr.find(w_namespace + 'docGrid')
                if docGrid is None:
                    # Fix the missing nsmap parameter by using an empty dict for attrib and None for nsmap
                    docGrid = etree.SubElement(sectPr, w_namespace + 'docGrid', {}, None)
                docGrid.set(w_namespace + 'type', 'lines')
                docGrid.set(w_namespace + 'linePitch', '360')  # 360 twips standard
            
            # Write the modified XML back
            tree.write('word/document.xml', xml_declaration=True, encoding='UTF-8')
            
        # Create a new docx with the modified XML
        with zipfile.ZipFile(output_docx_file, 'w') as outzip:
            # Copy all the files from original docx
            with zipfile.ZipFile(temp_docx, 'r') as inzip:
                for item in inzip.infolist():
                    if item.filename != 'word/document.xml':
                        outzip.writestr(item, inzip.read(item.filename))
            
            # Add our modified document.xml
            outzip.write('word/document.xml', 'word/document.xml')
        
        # Clean up temporary files
        try:
            os.remove(temp_docx)
            os.remove('word/document.xml')
            os.rmdir('word')
        except:
            pass
        
        print(f"✓ Tategaki document successfully created: {output_docx_file}")
        print("  - The document is formatted for vertical Japanese text")
        print("  - Standard B6 novel size (128mm × 182mm)")
        print("  - 'MS Mincho' font for traditional Japanese novel appearance")
        print("  - Proper paragraph and line spacing for fiction")
        print("")
        print("Note: The document should display properly in Microsoft Word.")
        print("      If the text isn't vertical, go to Page Layout > Text Direction > Rotate all text 90°")
        
        return output_docx_file
    except Exception as e:
        print(f"Error creating tategaki document: {e}")
        return None

def load_checkpoint():
    """Load checkpoint data if it exists"""
    if os.path.exists(CHECKPOINT_FILE):
        try:
            with open(CHECKPOINT_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except (json.JSONDecodeError, IOError) as e:
            print(f"Warning: Failed to load checkpoint file: {e}")
            return None
    return None

def verify_output_file():
    """Check if output file exists and is accessible"""
    try:
        if os.path.exists(OUTPUT_FILE):
            with open(OUTPUT_FILE, 'a', encoding='utf-8') as f:
                # Test if file is writable
                f.write("")
                f.flush()
                os.fsync(f.fileno())
            return True
    except IOError as e:
        print(f"Warning: Output file issue: {e}")
        return False
    return True

def save_to_file(text, append=True):
    """Save text to file with proper flushing to ensure it's written immediately"""
    try:
        mode = "a" if append else "w"
        with open(OUTPUT_FILE, mode, encoding="utf-8") as f:
            f.write(text)
            f.flush()
            os.fsync(f.fileno())  # Force write to disk
        return True
    except IOError as e:
        print(f"Error writing to file: {e}")
        return False

def configure_api():
    """Configure API settings through interactive prompts"""
    global HF_API_URL, HF_MODEL, HF_TOKEN, HEADERS
    
    print("Current API Configuration:")
    print(f"1. API URL: {HF_API_URL}")
    print(f"2. Model: {HF_MODEL}")
    print(f"3. API Token: {'[SET]' if HF_TOKEN else '[NOT SET]'}")
    print("4. Continue with these settings")
    print("5. Test current API connection")
    
    choice = input("Enter your choice (1-5): ").strip()
    
    if choice == "1":
        new_url = input(f"Enter new API URL [{HF_API_URL}]: ").strip()
        if new_url:
            HF_API_URL = new_url
        return configure_api()
    elif choice == "2":
        new_model = input(f"Enter model name [{HF_MODEL}]: ").strip()
        if new_model:
            HF_MODEL = new_model
        return configure_api()
    elif choice == "3":
        new_token = input("Enter API token: ").strip()
        if new_token:
            HF_TOKEN = new_token
            HEADERS = {"Authorization": f"Bearer {HF_TOKEN}"}
        return configure_api()
    elif choice == "5":
        test_api_connection()
        input("Press Enter to continue...")
        return configure_api()
    else:
        # Option 4 or any other input: continue with current settings
        return

def main():
    # Start the cloud backup thread
    start_backup_thread()
      # First, test the API connection
    print("Testing API connection before starting translation...")
    api_working = test_api_connection()
    
    if not api_working:
        print("\nWARNING: API connection test failed. You may encounter problems during translation.")
        print("Would you like to configure the API settings?")
        if input("Configure API? (y/n): ").strip().lower() == 'y':
            configure_api()
    else:
        print("\nAPI connection test successful!")
        
        # Test the template with a sample translation
        print("\nWould you like to test the translation template?")
        if input("Test template? (y/n): ").strip().lower() == 'y':
            test_success = test_translation_template()
            
            if not test_success and not TRANSLATE_TEMPLATE_En:
                print("\nTranslation test failed. Would you like to set a recommended template?")
                if input("Set recommended template? (y/n): ").strip().lower() == 'y':
                    # Set a recommended template that has demonstrated success
                    TRANSLATE_TEMPLATE_En = (
                        "[INST] You are a professional Japanese translator. Your task is to translate the following English text into natural, fluent Japanese.\n\n"
                        "Translation style: Young adult fantasy novel for vertical writing (tategaki).\n"
                        "Express subtle wonder, emotional depth, and a slightly melancholic tone.\n"
                        "Use appropriate Japanese punctuation (「」for quotes, 。for periods, etc).\n"
                        "Use natural Japanese expressions rather than literal translations.\n\n"
                        "IMPORTANT: Return ONLY the Japanese translation. No explanations or English text.\n\n"
                        "Text to translate: {text} [/INST]"
                    )
                    print("\nTemplate set. Testing again with new template...")
                    test_translation_template()
    
    # Check for existing checkpoint
    checkpoint = load_checkpoint()
    resume_translation = False
    
    if checkpoint and os.path.exists(OUTPUT_FILE):
        response = input(f"Found checkpoint from previous translation. Resume? (y/n): ").strip().lower()
        resume_translation = response == 'y'
    
    if resume_translation and checkpoint:
        docx_path = checkpoint['docx_path']
        print(f"Resuming translation of {docx_path}")
        
        # We'll keep the existing output file
        content_units_processed = checkpoint['content_units_processed']
        paragraph_number = checkpoint['paragraph_number']
        manuscript_title_count = checkpoint['manuscript_title_count']
        chapter_title_count = checkpoint['chapter_title_count']
        start_time_offset = checkpoint['elapsed_time']
        # We'll adjust the start time later
    else:
        docx_path = input("Enter the full path to your DOCX manuscript: ").strip()
        if not os.path.isfile(docx_path):
            print(f"File not found: {docx_path}")
            sys.exit(1)

        # Clear output file at start of new translation
        save_to_file("", append=False)
        
        content_units_processed = 0
        paragraph_number = 0
        manuscript_title_count = 0
        chapter_title_count = 0
        start_time_offset = 0

    doc = docx.Document(docx_path)    # First pass: Count all meaningful content units (sentences and titles)
    total_content_units = 0
    total_paragraphs = 0
    paragraph_positions = []  # Track positions of each paragraph for resuming
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue  # skip empty paragraphs
        
        if text.startswith("###") or text.startswith("##"):
            # Titles count as 1 content unit each
            total_content_units += 1
            paragraph_positions.append({
                'index': i, 
                'type': 'title', 
                'text': text
            })
        else:
            # Regular paragraphs: count sentences
            sentences = nltk.sent_tokenize(text)
            total_content_units += len(sentences)
            total_paragraphs += 1
            paragraph_positions.append({
                'index': i, 
                'type': 'paragraph', 
                'text': text, 
                'sentence_count': len(sentences)
            })

    print(f"Document contains {total_content_units} content units to translate ({total_paragraphs} paragraphs)")
    
    # Initialize counters based on whether we're resuming or not
    if not resume_translation:
        paragraph_number = 0
        manuscript_title_count = 0
        chapter_title_count = 0
        content_units_processed = 0
        start_time_offset = 0
    
    # Adjust start time based on elapsed time from checkpoint
    start_time = time.time() - start_time_offset
    
    para_idx = 0
    try:
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue  # skip empty paragraphs
            
            # If we're resuming and haven't reached the checkpoint position, skip this paragraph
            if resume_translation and checkpoint and para_idx < checkpoint.get('last_paragraph_index', 0):
                para_idx += 1
                continue

            if text.startswith("###"):  # Manuscript title line
                manuscript_title_count += 1
                to_translate = text[3:].strip()
                
                # Update progress
                content_units_processed += 1
                percent_done = (content_units_processed / total_content_units) * 100
                eta = calculate_eta(start_time, content_units_processed, total_content_units)
                
                print(f"Translating manuscript title {manuscript_title_count}: {to_translate}")
                
                try:
                    jp_translation = translate_text(to_translate)
                    print(f"Done: {content_units_processed}/{total_content_units} ({percent_done:.1f}%) ETA: {eta} [{jp_translation}]")
                      # Remove markdown formatting for output file
                    save_to_file(f"{jp_translation}\n\n")
                    
                    # Save checkpoint after each unit
                    checkpoint_data = {
                        'docx_path': docx_path,
                        'content_units_processed': content_units_processed,
                        'paragraph_number': paragraph_number,
                        'manuscript_title_count': manuscript_title_count,
                        'chapter_title_count': chapter_title_count,
                        'last_paragraph_index': para_idx,
                        'elapsed_time': time.time() - start_time
                    }
                    save_checkpoint(checkpoint_data)
                    
                except Exception as e:
                    print(f"Error translating manuscript title {manuscript_title_count}: {e}")
                    save_to_file(f"### [Translation error]\n\n")

            elif text.startswith("##"):  # Chapter title line
                chapter_title_count += 1
                to_translate = text[2:].strip()
                
                # Update progress
                content_units_processed += 1
                percent_done = (content_units_processed / total_content_units) * 100
                eta = calculate_eta(start_time, content_units_processed, total_content_units)
                
                print(f"Translating chapter title {chapter_title_count}: {to_translate}")
                
                try:
                    jp_translation = translate_text(to_translate)
                    print(f"Done: {content_units_processed}/{total_content_units} ({percent_done:.1f}%) ETA: {eta} [{jp_translation}]")
                      # Remove markdown formatting for output file
                    save_to_file(f"{jp_translation}\n\n")
                    
                    # Save checkpoint after each unit
                    checkpoint_data = {
                        'docx_path': docx_path,
                        'content_units_processed': content_units_processed,
                        'paragraph_number': paragraph_number,
                        'manuscript_title_count': manuscript_title_count, 
                        'chapter_title_count': chapter_title_count,
                        'last_paragraph_index': para_idx,
                        'elapsed_time': time.time() - start_time
                    }
                    save_checkpoint(checkpoint_data)
                    
                except Exception as e:
                    print(f"Error translating chapter title {chapter_title_count}: {e}")
                    save_to_file(f"## [Translation error]\n\n")

            else:
                # Only increment paragraph count for regular paragraphs
                paragraph_number += 1
                print(f"Processing paragraph {paragraph_number}...")
                
                # Regular paragraph: split into sentences
                sentences = nltk.sent_tokenize(text)
                
                # If resuming, check if we need to skip some sentences in the current paragraph
                sentence_start_idx = 0
                if resume_translation and checkpoint and para_idx == checkpoint.get('last_paragraph_index', 0):
                    # We might need to skip some sentences in this paragraph
                    sentence_start_idx = checkpoint.get('last_sentence_index', -1) + 1
                    
                for i, sentence in enumerate(sentences):
                    # Skip sentences we've already processed during resume
                    if i < sentence_start_idx:
                        continue
                        
                    # Update progress for each sentence
                    content_units_processed += 1
                    percent_done = (content_units_processed / total_content_units) * 100
                    eta = calculate_eta(start_time, content_units_processed, total_content_units)
                    
                    print(f"Translating sentence {i+1}/{len(sentences)} of paragraph {paragraph_number}...")
                    
                    try:
                        jp_translation = translate_text(sentence)
                        print(f"Done: {content_units_processed}/{total_content_units} ({percent_done:.1f}%) ETA: {eta} [{jp_translation}]")
                        
                        # Write each sentence immediately and flush to disk
                        save_to_file(jp_translation + "\n")
                        
                        # Save checkpoint after each sentence
                        checkpoint_data = {
                            'docx_path': docx_path,
                            'content_units_processed': content_units_processed,
                            'paragraph_number': paragraph_number,
                            'manuscript_title_count': manuscript_title_count,
                            'chapter_title_count': chapter_title_count,
                            'last_paragraph_index': para_idx,
                            'last_sentence_index': i,
                            'elapsed_time': time.time() - start_time
                        }
                        save_checkpoint(checkpoint_data)
                        
                    except Exception as e:
                        print(f"Error translating sentence {i+1} of paragraph {paragraph_number}: {e}")
                        save_to_file("[Translation error]\n")
                
                # Paragraph break
                save_to_file("\n")
            
            para_idx += 1
    except KeyboardInterrupt:
        # Handle user interruption gracefully
        print("\nTranslation paused. You can resume later.")
        elapsed = time.time() - start_time
        print(f"Progress: {content_units_processed}/{total_content_units} ({content_units_processed/total_content_units*100:.1f}%)")
        print(f"Time elapsed so far: {str(timedelta(seconds=int(elapsed)))}")
        
        # Perform backup before exiting
        if os.path.exists(OUTPUT_FILE):
            print("Performing backup to Google Cloud Storage before exit...")
            subprocess.run(["gsutil", "cp", OUTPUT_FILE, GCS_BUCKET])
            
        # Stop the backup thread before exiting
        stop_backup_thread()
        sys.exit(0)
    except Exception as e:
        # Handle any other exceptions
        print(f"\nError during translation: {e}")
        
        # Perform backup before exiting
        if os.path.exists(OUTPUT_FILE):
            print("Performing backup to Google Cloud Storage before exit...")
            subprocess.run(["gsutil", "cp", OUTPUT_FILE, GCS_BUCKET])
            
        # Stop the backup thread before exiting
        stop_backup_thread()
        raise
    
    # Translation complete - clean up checkpoint file
    if os.path.exists(CHECKPOINT_FILE):
        try:
            os.remove(CHECKPOINT_FILE)
            print("Checkpoint file removed as translation is complete.")
        except:
            pass
    
    total_time = time.time() - start_time
    hours, remainder = divmod(total_time, 3600)
    minutes, seconds = divmod(remainder, 60)
    
    # Perform one final backup before stopping the thread
    if os.path.exists(OUTPUT_FILE):
        print("Performing final backup to Google Cloud Storage...")
        subprocess.run(["gsutil", "cp", OUTPUT_FILE, GCS_BUCKET])
      # Stop the backup thread
    stop_backup_thread()
    
    print(f"\nTranslation complete! Output written to {OUTPUT_FILE}")
    print(f"Statistics: {manuscript_title_count} manuscript title(s), {chapter_title_count} chapter title(s), {paragraph_number} paragraphs")
    print(f"Total content units: {content_units_processed}/{total_content_units}")
    print(f"Total time taken: {int(hours)}h {int(minutes)}m {int(seconds)}s")

    # Create tategaki Word document
    create_tategaki_document(OUTPUT_FILE)

if __name__ == "__main__":
    main()
``` 
